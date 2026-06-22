from __future__ import annotations

import base64
import binascii
import io
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown"}


class DocumentIngestionError(Exception):
    def __init__(self, code: str, message: str, details: dict[str, Any] | None = None):
        super().__init__(message)
        self.code = code
        self.message = message
        self.details = details or {}

    def to_dict(self) -> dict[str, Any]:
        payload: dict[str, Any] = {"error": {"code": self.code, "message": self.message}}
        if self.details:
            payload["error"]["details"] = self.details
        return payload


@dataclass
class IngestedDocument:
    text: str
    metadata: dict[str, Any]


def ingest_document(*, filename: str, content_bytes: bytes) -> IngestedDocument:
    if not filename:
        raise DocumentIngestionError("file_missing", "File must include a filename")
    if content_bytes is None:
        raise DocumentIngestionError("file_missing", "File content is missing")
    if len(content_bytes) == 0:
        raise DocumentIngestionError("empty_document", "Document has no content")

    extension, parser_name = detect_file_type(filename, content_bytes)
    if parser_name == "pdf":
        text, page_count = parse_pdf(content_bytes)
    elif parser_name == "docx":
        text, page_count = parse_docx(content_bytes)
    else:
        text = parse_plain_text(content_bytes)
        page_count = None

    normalized_text = normalize_text(text)
    if not normalized_text:
        raise DocumentIngestionError("empty_document", "Document content is empty after parsing")

    metadata = {
        "filename": filename,
        "extension": extension,
        "page_count": page_count,
        "word_count": count_words(normalized_text),
        "character_count": len(normalized_text),
    }
    return IngestedDocument(text=normalized_text, metadata=metadata)


def ingest_document_payload(payload: dict[str, Any]) -> IngestedDocument:
    upload = payload.get("uploaded_file")
    if not isinstance(upload, dict):
        raise DocumentIngestionError("file_missing", "Missing uploaded_file payload")

    filename = str(upload.get("filename") or "").strip()
    encoded = upload.get("content_base64")
    if not isinstance(encoded, str) or not encoded.strip():
        raise DocumentIngestionError("file_missing", "uploaded_file.content_base64 is required")

    encoded_value = encoded.strip()
    if "," in encoded_value and "base64" in encoded_value[:40].lower():
        encoded_value = encoded_value.split(",", 1)[1].strip()

    try:
        content_bytes = base64.b64decode(encoded_value, validate=True)
    except (binascii.Error, ValueError):
        raise DocumentIngestionError("corrupted_document", "Invalid base64 content for uploaded file")

    return ingest_document(filename=filename, content_bytes=content_bytes)


def detect_file_type(filename: str, content_bytes: bytes) -> tuple[str, str]:
    extension = Path(filename).suffix.lower()

    if content_bytes.startswith(b"%PDF"):
        return ".pdf", "pdf"
    if extension == ".docx" or (content_bytes.startswith(b"PK") and looks_like_docx(content_bytes)):
        return ".docx", "docx"
    if extension in {".txt"}:
        return ".txt", "text"
    if extension in {".md", ".markdown"}:
        return extension, "text"

    if extension in SUPPORTED_EXTENSIONS:
        if extension == ".pdf":
            raise DocumentIngestionError("corrupted_document", "PDF signature not found")
        if extension == ".docx":
            raise DocumentIngestionError("corrupted_document", "DOCX structure is invalid")
        return extension, "text"

    raise DocumentIngestionError(
        "unsupported_file_type",
        "Unsupported document type",
        details={"filename": filename, "extension": extension or None, "supported": sorted(SUPPORTED_EXTENSIONS)},
    )


def looks_like_docx(content_bytes: bytes) -> bool:
    try:
        with zipfile.ZipFile(io.BytesIO(content_bytes)) as archive:
            names = set(archive.namelist())
            return "[Content_Types].xml" in names and "word/document.xml" in names
    except zipfile.BadZipFile:
        return False


def parse_pdf(content_bytes: bytes) -> tuple[str, int]:
    try:
        reader = PdfReader(io.BytesIO(content_bytes))
    except Exception as exc:
        raise DocumentIngestionError("corrupted_document", f"Unable to parse PDF: {exc}")

    fragments: list[str] = []
    for page in reader.pages:
        try:
            extracted = page.extract_text() or ""
        except Exception as exc:
            raise DocumentIngestionError("corrupted_document", f"Unable to extract PDF page text: {exc}")
        if extracted:
            fragments.append(extracted)

    return "\n".join(fragments), len(reader.pages) or None


def parse_docx(content_bytes: bytes) -> tuple[str, int | None]:
    try:
        document = Document(io.BytesIO(content_bytes))
    except Exception as exc:
        raise DocumentIngestionError("corrupted_document", f"Unable to parse DOCX: {exc}")

    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    return "\n".join(paragraphs), None


def parse_plain_text(content_bytes: bytes) -> str:
    try:
        return content_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return content_bytes.decode("latin-1")
        except UnicodeDecodeError as exc:
            raise DocumentIngestionError("corrupted_document", f"Unable to decode text document: {exc}")


def normalize_text(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"[\t\f\v]+", " ", normalized)
    normalized = "\n".join(line.strip() for line in normalized.split("\n"))
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def count_words(text: str) -> int:
    return len(re.findall(r"\b\w+\b", text, flags=re.UNICODE))
